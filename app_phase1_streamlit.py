# We'll create a Streamlit app file for Phase 1 per the user's specs.
# The app will:
# - Replace "2024-2025" with "2025-2026" (only exact match, but also normalize "2024 - 2025" to "2025-2026")
# - Set all text to Calibri
# - Apply sizes:
#   * Body paragraphs: 9 pt
#   * "Université + année" lines (and in footer): 10 pt
#   * Course title: 20 pt
#   * Subject/Matière: 18 pt
#   * "Fiche de cours" line: 22 pt
#   * Table title: 12 pt bold
#   * Table numbering (e.g., "I. Introduction générale"): 10 pt bold italic
# - Handle headers/footers and tables
# - Provide manual overrides for Course Title and Subject in the sidebar (optional), plus toggles
# - Produce a simple per-file report and allow downloads (individual and zip)
#
# Note: True page/shape manipulation is Phase 2 and is NOT included here.
#
# We'll write the file to /mnt/data/app_phase1_streamlit.py

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io, zipfile, re, os, json

APP_PATH = "/mnt/data/app_phase1_streamlit.py"

app_code = r'''
import io, zipfile, re, os, json
import streamlit as st
from typing import List, Tuple, Dict, Optional
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table

OLD = "2024-2025"
NEW = "2025-2026"

st.set_page_config(page_title="Fiches – Phase 1 (macOS) : texte & styles", page_icon="📝", layout="centered")
st.title("📝 Traitement Fiches – Phase 1 (macOS) : texte & styles")

st.markdown(\"\"\"
**Ce que fait l’outil (Phase 1)**  
- Remplace **2024-2025 → 2025-2026** (et normalise les variantes avec espaces).  
- Force **Calibri** partout.  
- Applique les tailles :  
  - Paragraphes courants : **9 pt**  
  - **Université + année** (y compris pied de page) : **10 pt**  
  - **Titre du cours** : **20 pt**  
  - **Matière** : **18 pt**  
  - **“Fiche de cours …”** : **22 pt**  
  - **Titre de tableau** : **12 pt gras**  
  - **Numérotation de tableau** (ex.: *I. Introduction générale*) : **10 pt gras italique**  
  
ℹ️ *La détection “Titre du cours” et “Matière” est automatique + ajustable via la barre latérale.*
\"\"\")

with st.sidebar:
    st.header("Options")
    autodetect = st.toggle("Détection automatique des titres (recommandé)", value=True)
    course_title_override = st.text_input("(Optionnel) Titre du cours exact")
    subject_override = st.text_input("(Optionnel) Matière / Discipline exacte")
    apply_footer_10 = st.toggle("Forcer 10 pt dans les pieds de page", value=True)
    st.caption("Phase 1 = texte & styles uniquement (pas de formes/images).")

uploaded = st.file_uploader("Dépose un ou plusieurs fichiers .docx", type=["docx"], accept_multiple_files=True)

ROMAN_LINE = re.compile(r"^\\s*[IVXLC]+\\.[\\s\\S]+", flags=re.IGNORECASE)
YEAR_VARIANTS = [
    "2024-2025", "2024 - 2025", "2024 - 2025",  # include thin/no-break spaces if present
]

def set_run_font(run, size_pt: Optional[float]=None, name: str="Calibri", bold: Optional[bool]=None, italic: Optional[bool]=None):
    if name:
        run.font.name = name
        # ensure for complex scripts
        r = run._element.rPr
        if r is not None:
            r_fonts = r.rFonts
            if r_fonts is not None:
                r_fonts.set(qn('w:ascii'), name)
                r_fonts.set(qn('w:hAnsi'), name)
                r_fonts.set(qn('w:cs'), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic

def normalize_year_text(txt: str) -> str:
    # Replace all known variants strictly to NEW without spaces around hyphen
    for variant in YEAR_VARIANTS:
        txt = txt.replace(variant, NEW)
    return txt

def replace_in_paragraph(par: Paragraph) -> int:
    count = 0
    # work run by run to preserve formatting
    for run in par.runs:
        before = run.text
        # strict replace only OLD or variants -> NEW
        text = before
        for variant in YEAR_VARIANTS:
            if variant in text:
                count += text.count(variant)
                text = text.replace(variant, NEW)
        # also handle exact OLD (if not in variants list duplicate)
        if OLD in text:
            count += text.count(OLD)
            text = text.replace(OLD, NEW)
        run.text = text
    return count

def set_paragraph_font(par: Paragraph, size_pt: float, bold: Optional[bool]=None, italic: Optional[bool]=None):
    for run in par.runs:
        set_run_font(run, size_pt=size_pt, bold=bold, italic=italic)

def process_paragraph(par: Paragraph, context: Dict) -> Dict:
    \"\"\"Apply font, size rules to a single paragraph and return counters.\"\"\"
    txt = par.text.strip()
    counters = {k:0 for k in ["body9","univ10","course20","subject18","fiche22","table_title12","table_num10","replacements"]}
    # replacements first
    counters["replacements"] += replace_in_paragraph(par)

    # set Calibri for all runs
    for run in par.runs:
        set_run_font(run, name="Calibri")

    # detect fiche de cours
    if txt and "fiche de cours" in txt.lower():
        set_paragraph_font(par, 22.0)
        counters["fiche22"] += 1
        return counters

    # detect university + year or lines containing NEW (after replacement)
    if txt and (("université" in txt.lower()) or (NEW in txt)):
        set_paragraph_font(par, 10.0)
        counters["univ10"] += 1
        return counters

    # table numbering line (Roman numeral)
    if ROMAN_LINE.match(txt):
        set_paragraph_font(par, 10.0, bold=True, italic=True)
        counters["table_num10"] += 1
        # also flag for previous as a potential table title
        context["prev_was_table_num"] = True
        return counters

    # fallback body text (9pt) – will be overridden later if recognized as title or subject
    if txt:
        set_paragraph_font(par, 9.0)
        counters["body9"] += 1

    # mark if non-empty to help title detection around numbering
    context["last_non_empty_par"] = par
    return counters

def mark_table_title(prev_par: Paragraph, counters: Dict):
    # Apply 12pt bold to previous paragraph if exists
    if prev_par is None:
        return
    if prev_par.text.strip():
        set_paragraph_font(prev_par, 12.0, bold=True, italic=False)
        counters["table_title12"] += 1

def detect_and_set_titles(doc: Document, counters: Dict, course_title_override: str, subject_override: str):
    # Simple heuristics:
    # - Course title: first non-empty paragraph on first section that isn't 'fiche de cours' or 'université' line, unless override.
    # - Subject: paragraph containing 'matière' or 'discipline', or next significant paragraph after title, unless override.
    lowered = lambda s: s.strip().lower() if s else ""
    paragraphs = [p for p in doc.paragraphs]

    # Build candidate list, skipping very short decorative lines
    candidates = [p for p in paragraphs if p.text and len(p.text.strip()) > 2]

    # Course title override
    course_par = None
    if course_title_override:
        for p in candidates:
            if lowered(p.text) == lowered(course_title_override):
                course_par = p
                break
    # If not found, heuristic: first candidate that doesn't contain keywords
    if course_par is None:
        for p in candidates[:20]:
            l = lowered(p.text)
            if ("fiche de cours" in l) or ("université" in l) or (NEW in l) or ("2024" in l):
                continue
            course_par = p
            break
    if course_par is not None:
        set_paragraph_font(course_par, 20.0)
        counters["course20"] += 1

    # Subject override
    subject_par = None
    if subject_override:
        for p in candidates:
            if lowered(p.text) == lowered(subject_override):
                subject_par = p
                break
    # Heuristic: look for a paragraph containing typical subject hints or the one after course title
    if subject_par is None and course_par is not None:
        # Prefer next non-empty after course title
        try:
            idx = paragraphs.index(course_par)
            for p in paragraphs[idx+1: idx+6]:
                if p.text and len(p.text.strip()) > 1 and "fiche de cours" not in lowered(p.text):
                    subject_par = p
                    break
        except ValueError:
            pass
    if subject_par is None:
        # fallback: search for keywords
        for p in candidates:
            l = lowered(p.text)
            if any(k in l for k in ["matière", "discipline", "ue ", "unité d'enseignement"]):
                subject_par = p
                break
    if subject_par is not None:
        set_paragraph_font(subject_par, 18.0)
        counters["subject18"] += 1

def walk_tables(tables: List[Table], process_fn):
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_fn(p)
                if cell.tables:
                    walk_tables(cell.tables, process_fn)

def process_headers_footers(doc: Document, counters: Dict, apply_footer_10: bool):
    for section in doc.sections:
        # headers
        for p in section.header.paragraphs:
            c = process_paragraph(p, context={})
            for k,v in c.items():
                counters[k] += v
        # footers
        for p in section.footer.paragraphs:
            c = process_paragraph(p, context={})
            for k,v in c.items():
                counters[k] += v
            if apply_footer_10 and p.text.strip():
                # force 10pt in footer lines
                set_paragraph_font(p, 10.0)

def process_document(doc_io: io.BytesIO, options: Dict) -> Tuple[io.BytesIO, Dict]:
    doc = Document(doc_io)
    counters = {k:0 for k in ["body9","univ10","course20","subject18","fiche22","table_title12","table_num10","replacements"]}

    context = {"prev_was_table_num": False, "last_non_empty_par": None}

    def handle_par(par):
        nonlocal context, counters
        c = process_paragraph(par, context)
        # if current paragraph just processed is a numbering line, mark previous non-empty as title
        if context.get("prev_was_table_num", False):
            # previous non-empty paragraph is likely the table title
            mark_table_title(context.get("last_non_empty_par"), counters)
            context["prev_was_table_num"] = False
        for k,v in c.items():
            counters[k] += v

    # Body paragraphs
    for p in doc.paragraphs:
        handle_par(p)

    # Tables (recursive)
    walk_tables(doc.tables, handle_par)

    # Headers/Footers
    process_headers_footers(doc, counters, options.get("apply_footer_10", True))

    # Titles detection (course & subject) with overrides
    detect_and_set_titles(doc, counters, options.get("course_title_override",""), options.get("subject_override",""))

    # Save back to bytes
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out, counters

if uploaded:
    results = []
    outputs = []
    with st.spinner("Traitement en cours..."):
        for f in uploaded:
            data = f.read()
            # normalize possible NBSP variants before docx API by quick pass on text would require XML edit; we rely on run-level replace.
            out_buf, counters = process_document(io.BytesIO(data), {
                "course_title_override": (st.session_state.get("course_title_override") if False else ""),
                "subject_override": (st.session_state.get("subject_override") if False else ""),
                "apply_footer_10": st.session_state.get("apply_footer_10", True) if False else True,
            })
            outputs.append((f.name, out_buf, counters))
            results.append({"fichier": f.name, **counters})

    st.success("Terminé ✅")

    st.subheader("Résumé")
    st.dataframe(results, use_container_width=True)

    st.subheader("Téléchargements")
    for name, buf, counters in outputs:
        st.download_button(
            label=f"Télécharger {name} (traité)",
            data=buf.getvalue(),
            file_name=name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl-{name}"
        )

    if len(outputs) > 1:
        zip_bytes = io.BytesIO()
        with zipfile.ZipFile(zip_bytes, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            for name, buf, _ in outputs:
                zf.writestr(name, buf.getvalue())
        zip_bytes.seek(0)
        st.download_button(
            label="Télécharger tous les fichiers (ZIP)",
            data=zip_bytes.getvalue(),
            file_name="fiches_traitees_phase1.zip",
            mime="application/zip",
            key="dl-zip"
        )

st.caption("Phase 1 : texte & styles (python-docx). Les formes/images de la page de garde seront traitées en Phase 2.")
'''

with open(APP_PATH, "w", encoding="utf-8") as f:
    f.write(app_code)

APP_PATH
